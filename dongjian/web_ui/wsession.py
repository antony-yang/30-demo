import datetime
import multiprocessing
import subprocess
import time

import psutil
import json
import traceback
import math
import socket
import os
import shutil
import copy

import requests

from web_ui import web_config

from gridfs import GridFSBucket
from scapy.sendrecv import sendp
from scapy.utils import rdpcap

import time
from DongJian.utils import CommonDBhandler
from DongJian.utils.guan_http import GuanHTTP
import threading
from DongJian.utils import db_config

from docx import Document
import xml.dom.minidom
import xlwt

import logging

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

protos = {}
protos2 = []
for parent, _, filenames in os.walk("../script"):
    for name in filenames:
        if name.startswith("."):
            continue
        if name.endswith("py"):
            script_obj = None
            script_name = name.split('.')[0]
            code = "from script import " + script_name + " as script_obj"
            exec(code)
            protos[script_name] = script_obj
            protos2.append(script_name)
        else:
            continue


class Wsession(object):
    def __init__(self):
        self._processes = {}
        self.processes = []
        self._db_handler = CommonDBhandler.handler

        # mongodb
        self._m_db = self._db_handler.getMC()["dongjian"]
        self._m_db_collection_cases = None
        self._m_db_collection_steps = None
        self._m_db_collection_crash_bin = None

        self._pause = 0
        self.result_total_page = 0
        self.result_data_page = 0
        self.result_crash_page = 0
        self.queue = multiprocessing.Queue(6)
        t = threading.Thread(target=self.process_end)
        t.start()
        self.func_desc = ""
        self.debug = True

    def gen_doc(self, run_id):
        collection_cases = self._m_db[run_id + "-cases"]
        collection_steps = self._m_db[run_id + "-steps"]
        collection_crash_bin = self._m_db[run_id + "-crash_bin"]
        task_list = self._m_db["task_list"]
        mclient = self._db_handler.getMC()
        mdb = mclient["report"]
        report = GridFSBucket(mdb, bucket_name=run_id + "-report")

        task_found = task_list.find({"run_id": run_id}, {"_id": 0})
        task = task_found.next()
        parameter = task["params"]
        taskname = task["taskname"]

        crash_info_find = collection_crash_bin.find({})
        crash_info = ""
        for crash in crash_info_find:
            crash_info += str(crash)

        pport = parameter['pport']
        proc_name = parameter['proc_name']
        protocol = parameter['protocol']
        target_ip = parameter['target_ip']
        taskcreator = parameter['taskcreator']
        tasktype = parameter['tasktype']
        start_cmds = parameter['start_cmds']

        total_cases = "N/A"
        tested_cases = str(collection_cases.count_documents({}))
        cases_list = collection_cases.find()

        start_time = task["ctime"]
        end_time = task["etime"]
        time_consuming = str((datetime.datetime.strptime(end_time, "%Y-%m-%dT%H-%M-%S-%f") - datetime.datetime.strptime(
            start_time, "%Y-%m-%dT%H-%M-%S-%f")).seconds) + "秒"
        task_status = task["status"]

        document = Document()
        document.add_heading(u'任务' + run_id + u'报告', 0)
        document.add_heading(u'一、任务摘要', level=1)

        document.add_paragraph(
            u'任务名称：' + taskname + "-" + run_id, style='ListNumber'
        )

        document.add_paragraph(
            u'任务创建时间:' + start_time, style='ListNumber'
        )

        document.add_paragraph(
            u'fuzz协议：' + protocol, style='ListNumber'
        )

        document.add_paragraph(
            u'协议名称：' + proc_name, style='ListNumber'
        )

        document.add_paragraph(
            u'主机ip：' + target_ip, style='ListNumber'
        )

        document.add_paragraph(
            u'端口号：' + str(pport), style='ListNumber'
        )

        document.add_paragraph(
            u'fuzz参数：' + ''.join(start_cmds), style='ListNumber'
        )

        document.add_heading(u'二、任务信息', level=1)

        document.add_paragraph(
            u'总测试用例数量:' + total_cases, style='ListNumber'
        )
        document.add_paragraph(
            u'已测测试用例数量:' + tested_cases, style='ListNumber'
        )
        document.add_paragraph(
            u'开始时间：' + start_time, style='ListNumber'
        )
        document.add_paragraph(
            u'结束时间：' + end_time, style='ListNumber'
        )
        document.add_paragraph(
            u'总耗时：' + time_consuming, style='ListNumber'
        )
        document.add_paragraph(
            u'任务状态信息：' + task_status, style='ListNumber'
        )

        document.add_heading(u'三、测试用例', level=1)
        for case in cases_list:
            document.add_paragraph(
                str(case), style='ListNumber'
            )

        document.add_heading(u'四、崩溃点信息', level=1)
        for crash in crash_info_find:
            document.add_paragraph(
                str(crash), style='ListNumber'
            )

        document.add_page_break()
        if os.path.exists("../doc"):
            pass
        else:
            os.mkdir("../doc")
        document.save('../doc/' + run_id + ".doc")
        fp = open(r'../doc/' + run_id + '.doc', 'rb')
        report.upload_from_stream(run_id + ".doc", fp)
        fp.close()

        # pdf
        if os.path.exists("../pdf"):
            pass
        else:
            os.mkdir("../pdf")
        base_path = os.getcwd().rsplit("/", 1)[0]
        cmd = 'libreoffice --headless --convert-to pdf ' + base_path + '/doc/' + run_id + ".doc --outdir " + base_path + "/pdf/"
        p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE, shell=True)
        p.wait(timeout=30)
        stdout, stderr = p.communicate()
        # time.sleep(2)
        fp = open(r'../pdf/' + run_id + '.pdf', 'rb')
        report.upload_from_stream(run_id + ".pdf", fp)
        fp.close()

        # xml
        doc = xml.dom.minidom.Document()
        root = doc.createElement('report')
        # 给根节点添加属性
        root.setAttribute('report', u'任务' + run_id + u'报告')
        # 将根节点添加到文档对象中
        doc.appendChild(root)
        summary = doc.createElement('summary')

        taskname_x = doc.createElement('taskname')
        taskname_x.appendChild(doc.createTextNode(run_id))

        creat_time_x = doc.createElement('creat_time')
        creat_time_x.appendChild(doc.createTextNode(start_time))

        fuzzing_agree = doc.createElement('protocol')
        fuzzing_agree.appendChild(doc.createTextNode(protocol))

        host_ip = doc.createElement('host_ip')
        host_ip.appendChild(doc.createTextNode(target_ip))

        port_x = doc.createElement('port')
        port_x.appendChild(doc.createTextNode(str(pport)))

        fuzzing_param = doc.createElement('fuzzing_param')
        fuzzing_param.appendChild(doc.createTextNode(''.join(start_cmds)))

        summary.appendChild(taskname_x)
        summary.appendChild(creat_time_x)
        summary.appendChild(fuzzing_agree)
        summary.appendChild(host_ip)
        summary.appendChild(port_x)
        summary.appendChild(fuzzing_param)

        task_info = doc.createElement('task_info')

        total_cases_x = doc.createElement('total_cases')
        total_cases_x.appendChild(doc.createTextNode(total_cases))

        tasted_cases_x = doc.createElement('tested_cases')
        tasted_cases_x.appendChild(doc.createTextNode(tested_cases))

        start_time_x = doc.createElement('start_time')
        start_time_x.appendChild(doc.createTextNode(start_time))

        end_time_x = doc.createElement('end_time')
        end_time_x.appendChild(doc.createTextNode(str(end_time)))

        time_consuming_x = doc.createElement('time_consuming')
        time_consuming_x.appendChild(doc.createTextNode(str(time_consuming)))

        task_status_x = doc.createElement('task_status')
        task_status_x.appendChild(doc.createTextNode(str(task_status)))

        task_info.appendChild(total_cases_x)
        task_info.appendChild(tasted_cases_x)
        task_info.appendChild(start_time_x)
        task_info.appendChild(end_time_x)
        task_info.appendChild(time_consuming_x)
        task_info.appendChild(task_status_x)

        crash = doc.createElement('crash')

        crash_node = doc.createElement('crash_node')
        crash_node.appendChild(doc.createTextNode(crash_info))
        crash.appendChild(crash_node)

        root.appendChild(summary)
        root.appendChild(task_info)
        root.appendChild(crash)

        if os.path.exists("../xml"):
            pass
        else:
            os.mkdir("../xml")
        fp = open(r'../xml/' + run_id + '.xml', 'w')
        doc.writexml(fp, indent='', addindent='\t', newl='\n', encoding='utf-8')
        fp.close()
        fp2 = open(r'../xml/' + run_id + '.xml', 'rb')
        report.upload_from_stream(run_id + ".xml", fp2)
        fp2.close()

        # excel
        report_excel = xlwt.Workbook(encoding='utf-8')

        summary_sheet = report_excel.add_sheet('一、任务摘要')

        summary_sheet.write(0, 0, '任务名称')
        summary_sheet.write(0, 1, taskname + "-" + run_id)

        summary_sheet.write(1, 0, u'任务创建时间')
        summary_sheet.write(1, 1, str(start_time))

        summary_sheet.write(2, 0, u'fuzz协议')
        summary_sheet.write(2, 1, protocol)

        summary_sheet.write(3, 0, u'协议名称')
        summary_sheet.write(3, 1, proc_name)

        summary_sheet.write(4, 0, u'主机ip')
        summary_sheet.write(4, 1, target_ip)

        summary_sheet.write(5, 0, u'端口号')
        summary_sheet.write(5, 1, pport)

        summary_sheet.write(6, 0, u'fuzz参数')
        summary_sheet.write(6, 1, start_cmds)

        info_sheet = report_excel.add_sheet(u'二、任务信息')

        info_sheet.write(0, 0, u'总测试用例数量')
        info_sheet.write(0, 1, str(total_cases))

        info_sheet.write(1, 0, u'已测测试用例数量')
        info_sheet.write(1, 1, str(tested_cases))

        info_sheet.write(2, 0, u'开始时间')
        info_sheet.write(2, 1, str(start_time))

        info_sheet.write(3, 0, u'结束时间')
        info_sheet.write(3, 1, str(end_time))

        info_sheet.write(4, 0, u'总耗时')
        info_sheet.write(4, 1, str(time_consuming))

        info_sheet.write(5, 0, u'任务状态信息')
        info_sheet.write(5, 1, str(task_status))

        crash_sheet = report_excel.add_sheet(u'三、崩溃点信息')
        crash_sheet.write(0, 0, crash_info)

        if os.path.exists("../xls"):
            pass
        else:
            os.mkdir("../xls")
        report_excel.save(r'../xls/' + run_id + '.xls')

        fp = open(r'../xls/' + run_id + '.xls', 'rb')
        report.upload_from_stream(run_id + ".xls", fp)
        fp.close()

    def get_all_protocol(self):
        return {"res": 1, "value": protos2}

    def get_deault_task_params(self, jsonp):
        return {"res": 1, "value": protos[jsonp["script_name"]].param}

    def create_task(self, p):
        try:
            p['protocol']
            p['target_ip']
            p['pport']
            p['taskname']
            p['tasktype']
            p['taskcreator']
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}
        pport = p['pport']
        if not (type(pport) == int):
            return {"res": 0, "value": "pport must be integers"}
        if pport > 65535 or pport < 0:
            return {"res": 0, "value": "pport should range from 0 to 65535"}
        run_id = datetime.datetime.now().isoformat().replace(":", "-").replace(".", "-")
        mrecord = {}
        mrecord["taskname"] = p["taskname"]
        mrecord["tasktype"] = p["tasktype"]
        mrecord["taskcreator"] = p["taskcreator"]
        mrecord["protocol"] = p["protocol"]
        mrecord["ctime"] = run_id
        mrecord["etime"] = ""
        mrecord["port"] = p["pport"]
        mrecord["targetip"] = p["target_ip"]
        mrecord["status"] = "unfinished"
        mrecord["run_id"] = run_id
        mrecord["params"] = p
        rmrecord = copy.deepcopy(mrecord)
        self._m_db["task_list"].insert_one(mrecord)
        return {"res": 1, "value": rmrecord}

    def insert_log(self, task_id, status, operation, information):
        log = self._m_db["log"]
        log_info = {}
        log_info["status"] = status
        log_info["category"] = operation
        log_info["detail"] = task_id + ":" + information
        log_info["date"] = int(round(time.time()*1000))
        log_info["time"] = datetime.datetime.now()
        log_info["name"] = "网络设备安全性测试系统"
        log.insert_one(log_info)

    def create_and_start_task(self, p):
        try:
            p['protocol']
            p['target_ip']
            p['pport']
            p['taskname']
            p['tasktype']
            p['taskcreator']
        except Exception as e:
            self.print_trace_back()
            return {"code": 400, "msg": "Missing " + str(e), "token": web_config.token}
        pport = p['pport']
        if not (type(pport) == int):
            return {"code": 400, "msg": "pport must be integers", "token": web_config.token}
        if pport > 65535 or pport < 0:
            return {"code": 400, "msg": "pport should range from 0 to 65535", "token": web_config.token}
        run_id = datetime.datetime.now().isoformat().replace(":", "-").replace(".", "-")
        mrecord = {}
        mrecord["taskname"] = p["taskname"]
        mrecord["tasktype"] = p["tasktype"]
        mrecord["taskcreator"] = p["taskcreator"]
        mrecord["protocol"] = p["protocol"]
        mrecord["ctime"] = run_id
        mrecord["etime"] = ""
        mrecord["port"] = p["pport"]
        mrecord["targetip"] = p["target_ip"]
        mrecord["status"] = "unfinished"
        mrecord["run_id"] = run_id
        mrecord["params"] = p
        mrecord["status"] = "unstart"
        mrecord["old_status"] = "unstart"
        rmrecord = copy.deepcopy(mrecord)
        if len(self.processes) >= 5:
            self.insert_log(status=False, task_id="", operation="create and start task",
                            information="The number of running tasks is >= 5, please wait.")
            return {"code": 400, "msg": "The number of running tasks is >= 5, please wait.", "token": web_config.token}
        self._m_db["task_list"].insert_one(mrecord)
        if self.start_task({"run_id": run_id})["res"] == 1:
            self.insert_log(status=True, task_id=run_id, operation="create and start task", information="success")
            return {"code": 200, "msg": "success", "token": web_config.token, "task_id": run_id}
        else:
            self.insert_log(status=False, task_id=run_id, operation="create and start task",
                            information="The task fails to start for unknown reason!")
            return {"code": 400, "msg": "The task fails to start for unknown reason!", "token": web_config.token}

    def start_task(self, p):
        if len(self.processes) >= 5:
            return {"res": 0, "value": "tasks > 5, not allowed!"}
        run_id = p["run_id"]
        if run_id in self._processes.keys():
            return {"res": 0, "value": "tasks already started!"}
        task_list_db = self._m_db["task_list"]
        task = task_list_db.find_one({"run_id": run_id}, {"_id": 0})
        kargs = task["params"]
        kargs["queue"] = self.queue
        kargs["run_id"] = run_id
        protocol = kargs["protocol"]
        process = multiprocessing.Process(target=protos[protocol].fuzz, kwargs=kargs, name=run_id)
        try:
            process.start()
            time.sleep(1)
            if process.is_alive():
                self._processes[run_id] = [str(process), "running"]
                self.processes.append(process)
                old_status = task["status"]
                self.push_status_change(run_id, "running", old_status)
                task_list_db.update_one({"run_id": run_id}, {'$set': {"status": "running", "old_status": old_status}})
                return {"res": 1, "value": "success"}
            else:
                return {"res": 0, "value": "task failed for unknown reason"}
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}

    def process_end(self):
        while True:
            if not self.queue.empty():
                run_id = self.queue.get(True)
                for process in self.processes:
                    if process.name == run_id:
                        self.processes.remove(process)
                        self._processes.pop(run_id)
                        break
                etime = datetime.datetime.now().isoformat().replace(":", "-").replace(".", "-")
                task_list_db = self._m_db["task_list"]
                task = task_list_db.find_one({"run_id": run_id}, {"_id": 0})
                old_status = task["status"]
                self.push_status_change(run_id, "finished", old_status)
                task_list_db.update_one({"run_id": run_id},
                                        {'$set': {"status": "finished", "etime": etime, "old_status": old_status}})
                self.gen_doc(run_id)
                self.push_task_result1(run_id)
                self.push_task_result2(run_id)
                self.push_task_result3(run_id)

    def get_running_task(self):
        task_list_db = self._m_db["task_list"]
        for process in self.processes:
            if not process.is_alive():
                self.processes.remove(process)
                self._processes.pop(process.name)
                task = task_list_db.find_one({"run_id": process.name}, {"_id": 0})
                old_status = task["status"]
                self.push_status_change(process.name, "finished", old_status)
                task_list_db.update_one({"run_id": process.name},
                                        {'$set': {"status": "finished", "old_status": old_status}})
        return {"res": 1, "value": self._processes}

    def kill_all_running_task(self):
        task_list_db = self._m_db["task_list"]
        for process in self.processes:
            process.terminate()
            etime = datetime.datetime.now().isoformat().replace(":", "-").replace(".", "-")
            task = task_list_db.find_one({"run_id": process.name}, {"_id": 0})
            old_status = task["status"]
            self.push_status_change(process.name, "finished", old_status)
            task_list_db.update_one({"run_id": process.name},
                                    {'$set': {"status": "finished", "etime": etime, "old_status": old_status}})
            # self.gen_doc(process.name)
        self._processes = {}
        self.processes = []
        return {"res": 1, "value": "success"}

    def kill_task_by_name(self, p):
        task_list_db = self._m_db["task_list"]
        try:
            p['run_id']
        except Exception as e:
            self.print_trace_back()
            print(1111)
            return {"res": 0, "value": str(e)}
        run_id = p['run_id']
        for process in self.processes:
            if process.name == run_id:
                self.processes.remove(process)
                self._processes.pop(run_id)
                process.terminate()
                etime = datetime.datetime.now().isoformat().replace(":", "-").replace(".", "-")
                task = task_list_db.find_one({"run_id": process.name}, {"_id": 0})
                old_status = task["status"]
                self.push_status_change(run_id, "finished", old_status)
                task_list_db.update_one({"run_id": run_id},
                                        {'$set': {"status": "finished", "etime": etime, "old_status": old_status}})
                #  self.gen_doc(run_id)
                #  self.push_task_result1(run_id)
                #  self.push_task_result2(run_id)
                #  self.push_task_result3(run_id)
                return {"res": 1, "value": "success"}
        return {"res": 0, "value": "no such task"}

    def suspend_all_task(self):
        task_list_db = self._m_db["task_list"]
        for process in self.processes:
            if self._processes[process.name][1] == "running":
                p = psutil.Process(process.pid)
                p.suspend()
                self._processes[process.name][1] = "paused"
                self.push_status_change(process.name, "paused", "running")
                task_list_db.update_one({"run_id": process.name},
                                        {'$set': {"status": "paused", "old_status": "running"}})
        return {"res": 1, "value": "success"}

    def suspend_task_by_name(self, p):
        task_list_db = self._m_db["task_list"]
        try:
            p['run_id']
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}
        run_id = p['run_id']
        for process in self.processes:
            if process.name == run_id:
                if self._processes[run_id][1] == "running":
                    p = psutil.Process(process.pid)
                    p.suspend()
                    self._processes[run_id][1] = "paused"
                    # self.push_status_change(run_id, "paused", "running")
                    task_list_db.update_one({"run_id": run_id}, {'$set': {"status": "paused", "old_status": "running"}})
                    return {"res": 1, "value": "success"}
        return {"res": 0, "value": "no such task"}

    def resume_all_task(self):
        task_list_db = self._m_db["task_list"]
        for process in self.processes:
            if self._processes[process.name][1] == "paused":
                p = psutil.Process(process.pid)
                p.resume()
                self._processes[process.name][1] = "running"
                self.push_status_change(process.name, "running", "paused")
                task_list_db.update_one({"run_id": process.name},
                                        {'$set': {"status": "running", "old_status": "paused"}})
        return {"res": 1, "value": "success"}

    def resume_task_by_name(self, p):
        task_list_db = self._m_db["task_list"]
        try:
            p['run_id']
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}
        run_id = p['run_id']
        flag = 0
        for process in self.processes:
            if process.name == run_id:
                if self._processes[run_id][1] == "paused":
                    flag = 1
                    p = psutil.Process(process.pid)
                    p.resume()
                    self._processes[run_id][1] = "running"
                    self.push_status_change(run_id, "running", "paused")
                    task_list_db.update_one({"run_id": run_id}, {'$set': {"status": "running", "old_status": "paused"}})
                    break
        if flag == 0:
            return {"res": 0, "value": "no such task"}
        return {"res": 1, "value": "success"}

    def control(self, p):
        try:
            p['taskIds']
            p['sig']
        except Exception as e:
            self.print_trace_back()
            self.insert_log(status=False, task_id="", operation="control task",
                            information="Lacking parameter " + str(e))
            return {"code": 400, "msg": "Lacking parameter " + str(e), "token": web_config.token}
        sig = p['sig']
        flag = 1
        if sig == "01":
            for taskid in p['taskIds']:
                if self.suspend_task_by_name({"run_id": taskid})["res"] == 0:
                    flag = 0
                    self.insert_log(status=False, task_id=taskid, operation="suspend task",
                                    information="Failed to suspend task.")
                self.insert_log(status=True, task_id=taskid, operation="suspend task", information="success.")
            if flag == 1:
                return {"code": 200, "msg": "success", "token": web_config.token}
            return {"code": 400, "msg": "Failed to suspend some tasks.", "token": web_config.token}
        if sig == "02":
            for taskid in p['taskIds']:
                if self.resume_task_by_name({"run_id": taskid})["res"] == 0:
                    flag = 0
                    self.insert_log(status=False, task_id=taskid, operation="resume task",
                                    information="Failed to resume task.")
                self.insert_log(status=True, task_id=taskid, operation="resume task", information="success.")
            if flag == 1:
                return {"code": 200, "msg": "success", "token": web_config.token}
            return {"code": 400, "msg": "Failed to resume some tasks.", "token": web_config.token}
        if sig == "03":
            for taskid in p['taskIds']:
                if self.kill_task_by_name({"run_id": taskid})["res"] == 0:
                    flag = 0
                    self.insert_log(status=False, task_id=taskid, operation="kill task",
                                    information="Failed to kill task.")
                self.insert_log(status=True, task_id=taskid, operation="kill task", information="success.")
            if flag == 1:
                return {"code": 200, "msg": "success", "token": web_config.token}
            return {"code": 400, "msg": "Failed to stop some tasks.", "token": web_config.token}
        self.insert_log(status=False, task_id="", operation="control task", information="Wrong signal!")
        return {"code": 400, "msg": "Wrong signal!", "token": web_config.token}

    def status(self, p):
        try:
            p['taskIds']
        except Exception as e:
            self.print_trace_back()
            self.insert_log(status=False, task_id="", operation="get status", information="Lacking parameter " + str(e))
            return {"code": 400, "msg": "Lacking parameter " + str(e), "token": web_config.token}
        status_dic = {"unstart": "01", "running": "02", "waiting": "03", "pasused": "04", "finished": "06"}
        status = []
        old_status = []
        try:
            task_list_db = self._m_db["task_list"]
            for taskid in p['taskIds']:
                task = task_list_db.find_one({"run_id": taskid}, {"_id": 0})
                status.append(status_dic[task["status"]])
                old_status.append(status_dic[task["old_status"]])
                self.insert_log(status=True, task_id=taskid, operation="get status", information="success")
            return {"code": 200, "msg": "success", "token": web_config.token, "status": status,
                    "old_status": old_status}
        except Exception as e:
            self.print_trace_back()
            self.insert_log(status=False, task_id="", operation="get status", information=str(e))
            return {"code": 400, "msg": str(e), "token": web_config.token}

    def progress(self, p):
        try:
            p['taskIds']
        except Exception as e:
            self.print_trace_back()
            self.insert_log(status=False, task_id="", operation="get progress", information="Lacking parameter " + str(e))
            return {"code": 400, "msg": "Lacking parameter " + str(e), "token": web_config.token}
        percent = []
        try:
            task_list_db = self._m_db["task_list"]
            data = []
            for taskid in p['taskIds']:
                task = task_list_db.find_one({"run_id": taskid}, {"_id": 0})
                if task["status"] == "unstart":
                    data.append({'taskId': taskid, 'percent': 0})
                if task["status"] == "running":
                    data.append({'taskId': taskid, 'percent': 50})
                if task["status"] == "pasused":
                    data.append({'taskId': taskid, 'percent': 50})
                if task["status"] == "finished":
                    data.append({'taskId': taskid, 'percent': 100})
                self.insert_log(status=True, task_id=taskid, operation="get progress", information="success")
            return {"code": 200, "msg": "success", "token": web_config.token, "data": data}
        except Exception as e:
            self.print_trace_back()
            self.insert_log(status=False, task_id="", operation="get progress", information=str(e))
            return {"code": 400, "msg": str(e)}

    def get_task_list(self):
        task_list_db = self._m_db["task_list"]
        try:
            task_list = task_list_db.find({}, {"_id": 0})
            res = []
            for item in task_list:
                if len(item) == 0:
                    break
                else:
                    item.pop("params")
                    res.append(item)
            return {"res": 1, "value": res}
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}

    def get_task_list_by_page(self, page_id):
        task_list_db = self._m_db["task_list"]
        page_size = 11
        res = []
        try:
            c = task_list_db.count_documents({})
            page_max = int(c / page_size) + 1
            if page_id <= 0:
                return {"res": 0, "value": res}
            if page_id < page_max:
                ret = task_list_db.find({}, {"_id": 0}).limit(page_size).skip((page_id - 1) * page_size)
            else:
                ret = task_list_db.find({}, {"_id": 0}).limit(c - (page_max - 1) * page_size).skip(
                    (page_max - 1) * page_size)
            for item in ret:
                if len(item) == 0:
                    break
                else:
                    item.pop("params")
                    res.append(item)
            return {"res": 1, "value": res}
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}

    def get_task_list_total_page(self):
        task_list_db = self._m_db["task_list"]
        page_size = 11
        try:
            c = task_list_db.count_documents({})
            res = int(c / page_size) + 1
            # return {"res": 1, "value": res}
            return {"res": 1, "value": c}
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}

    def ModString(self, s):
        ret = ""
        for i in range(0, len(s) - 1, 2):
            ret += "\\" + "x" + s[i:i + 2]
        return ret

    def get_task_result(self, p):
        task_list_db = self._m_db["task_list"]
        ret = []
        try:
            p['run_id']
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}
        run_id = p['run_id']
        try:
            if task_list_db.count_documents({"run_id": run_id}) > 0:
                pass
            else:
                return {"res": 0, "value": "no such task"}
            self._m_db_collection_steps = self._m_db[run_id + "-steps"]
            ret_all = self._m_db_collection_steps.find({}, {"_id": 0, "test_case_index": 1, "type": 1, "description": 1,
                                                            "data": 1, "timestamp2": 1, "istrancated": 1}).limit(24)
            res_all = []
            for item in ret_all:
                if len(item) == 0:
                    pass
                else:
                    res_all.append(item)
            if len(res_all) != 0:
                res_list = []
                flag = res_all[0]["test_case_index"]
                for i in range(0, 24):
                    row = res_all[i]
                    if flag == row["test_case_index"]:
                        res_list.append(row)
                    else:
                        ret.append(res_list)
                        flag = row["test_case_index"]
                        res_list = []
                        res_list.append(row)
            return {"res": 1, "value": ret}
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}

    def type1_result(self, run_id):
        task_list_db = self._m_db["task_list"]
        task = task_list_db.find_one({"run_id": run_id}, {"_id": 0})
        result = {}
        collection_cases = self._m_db[run_id + "-cases"]
        collection_steps = self._m_db[run_id + "-steps"]
        collection_crash_bin = self._m_db[run_id + "-crash_bin"]
        task_list = self._m_db["task_list"]

        task_found = task_list.find({"run_id": run_id}, {"_id": 0})
        task = task_found.next()
        parameter = task["params"]
        result["taskname"] = task["taskname"]
        result["pport"] = parameter['pport']
        result["proc_name"] = parameter['proc_name']
        result["protocol"] = parameter['protocol']
        result["target_ip"] = parameter['target_ip']
        result["taskcreator"] = parameter['taskcreator']
        result["tasktype"] = parameter['tasktype']
        result["start_cmds"] = parameter['start_cmds']
        start_time = task["ctime"]
        result["start_time"] = start_time
        end_time = task["etime"]
        result["end_time"] = end_time
        result["time_consuming"] = str(
            (datetime.datetime.strptime(end_time, "%Y-%m-%dT%H-%M-%S-%f") - datetime.datetime.strptime(
                start_time, "%Y-%m-%dT%H-%M-%S-%f")).seconds) + "秒"
        result["task_status"] = task["status"]

        crash_info_find = collection_crash_bin.find({}, {"_id": 0, "_len_": 0})
        crashes_list = []
        for crash in crash_info_find:
            crashes_list.append(crash)
        cases_info = collection_cases.find({}, {"_id": 0, "_len_": 0})
        cases_list = []
        for case in cases_info:
            cases_list.append(case)
        result["cases"] = cases_list
        result["crashes"] = crashes_list
        return result

    def push_task_result1(self, run_id):
        result = self.type1_result(run_id)
        args = {'taskId': run_id, 'token': web_config.token, "type": 1, 'resdata': result, "bianm": web_config.bianma}
        url = web_config.result_push_url
        request = GuanHTTP()
        try:
            response = request.post(url=url, json=args)
            if response["res"] == 0:
                self.insert_log(status=False, task_id="", operation="push type 1 result", information=str(response["value"]))
                logger.error(str(response["value"]))
            else:
                self.insert_log(status=True, task_id="", operation="push type 1 result", information="success")
                logger.info("push type 1 result success")
        except Exception as e:
            print("can't reach result_push_url")

    def get_task_result1(self, p):
        try:
            p['taskId']
        except Exception as e:
            self.insert_log(status=False, task_id="", operation="get task result, type 1", information="Lacking parameter " + str(e))
            return {"code": 400, "msg": "Lacking parameter " + str(e), "token": web_config.token}
        run_id = p['taskId']
        task_list_db = self._m_db["task_list"]
        task = task_list_db.find_one({"run_id": run_id}, {"_id": 0})
        if task["status"] != "finished":
            self.insert_log(status=False, task_id=run_id, operation="get task result, type 2",
                            information="The task is not fininshed.")
            return {"code": 400, "msg": "The task is not fininshed.", "token": web_config.token}
        result = self.type1_result(run_id)
        self.insert_log(status=True, task_id=run_id, operation="get task result, type 1", information="success")
        return {"code": 200, "msg": "success", "token": web_config.token, "resdata": result}

    def gen_zip(self, run_id):
        mclient = self._db_handler.getMC()
        mdb = mclient["report"]
        try:
            TMP_PATH = "../zip"
            if not os.path.exists(TMP_PATH + "/" + run_id):
                os.makedirs(TMP_PATH + "/" + run_id)
            buket = GridFSBucket(mdb, bucket_name=run_id + "-report")
            count = 0
            for file in buket.find():
                count = count + 1
                docfile = TMP_PATH + "/" + run_id + "/" + file.filename
                with open(docfile, "wb") as f:
                    buket.download_to_stream_by_name(file.filename, f)
            if count != 4:
                return 2
            compress_cmd = "cd {0};/bin/tar czf {1}.tar.gz {1}".format(TMP_PATH, run_id)
            os.system(compress_cmd)
            shutil.rmtree(TMP_PATH + "/" + run_id)
            logger.info("{} zip report generate success.".format(run_id))
            return 1
        except Exception as e:
            self.print_trace_back()
            logger.error(str(e))
            return 0

    def push_task_result2(self, run_id):
        filename = "../zip/" + run_id + ".tar.gz"
        if not os.path.exists(filename):
            if not self.gen_zip(run_id):
                logger.error("An error occurred while generating the zip package")
        file = {'file': open(filename, 'rb')}
        args = {'taskId': run_id, 'token': web_config.token, "type": 2, "bianm": web_config.bianma}
        url = web_config.result_push_url
        request = GuanHTTP()
        try:
            response = request.post(url=url, json=args, files=file)
            if response["res"] == 0:
                self.insert_log(status=False, task_id="", operation="push type 2 result", information=str(response["value"]))
                logger.error(str(response["value"]))
            else:
                self.insert_log(status=True, task_id="", operation="push type 2 result", information="success")
                logger.info("push type 2 result success")
        except Exception as e:
            print("can't reach result_push_url")

    def get_task_result2(self, p):
        try:
            p['taskId']
        except Exception as e:
            self.print_trace_back()
            self.insert_log(status=False, task_id="", operation="get task result, type 2", information="Lacking parameter " + str(e))
            return {"code": 400, "msg": "Lacking parameter " + str(e), "token": web_config.token}
        run_id = p['taskId']
        task_list_db = self._m_db["task_list"]
        task = task_list_db.find_one({"run_id": run_id}, {"_id": 0})
        if task["status"] != "finished":
            self.insert_log(status=False, task_id=run_id, operation="get task result, type 2",
                            information="The task is not fininshed.")
            return {"code": 400, "msg": "The task is not fininshed.", "token": web_config.token}
        filename = "../zip/" + run_id + ".tar.gz"
        if not os.path.exists(filename):
            if self.gen_zip(run_id) == 0:
                self.insert_log(status=False, task_id=run_id, operation="get task result, type 2",
                                information="An error occurred while generating the zip package.")
                return {"code": 400, "msg": "An error occurred while generating the zip package",
                        "token": web_config.token}
            if self.gen_zip(run_id) == 2:
                self.insert_log(status=False, task_id=run_id, operation="get task result, type 2",
                                information="The report has not been produced yet, please wait and try again")
                return {"code": 400, "msg": "The report has not been produced yet, please wait and try again",
                        "token": web_config.token}
        return run_id + ".tar.gz"

    def push_task_result3(self, run_id):
        args = {'taskId': run_id, 'token': web_config.token, "type": 3, "levels": None, "shuls": None, "types": None, "bianm": web_config.bianma}
        url = web_config.result_push_url
        request = GuanHTTP()
        try:
            response = request.post(url=url, json=args)
            if response["res"] == 0:
                self.insert_log(status=False, task_id="", operation="push type 3 result", information=str(response["value"]))
                logger.error(str(response["value"]))
            else:
                self.insert_log(status=True, task_id="", operation="push type 3 result", information="success")
                logger.info("push type 3 result success")
        except Exception as e:
            print("can't reach result_push_url")

    def get_task_result3(self, p):
        try:
            p['taskId']
        except Exception as e:
            self.print_trace_back()
            self.insert_log(status=False, task_id="", operation="get task result, type 3", information="Lacking parameter " + str(e))
            return {"code": 400, "msg": "Lacking parameter " + str(e), "token": web_config.token}
        run_id = p['taskId']
        task_list_db = self._m_db["task_list"]
        task = task_list_db.find_one({"run_id": run_id}, {"_id": 0})
        if task["status"] != "finished":
            self.insert_log(status=False, task_id=run_id, operation="get task result, type 3",
                            information="The task is not fininshed.")
            return {"code": 400, "msg": "The task is not fininshed.", "token": web_config.token}
        self.insert_log(status=True, task_id=run_id, operation="get task result, type 3", information="success.")
        return {"code": 200, "msg": "success", "token": web_config.token, "levels": None, "shuls": None, "types": None}

    def push_status_change(self, run_id, status, old_status):
        status_dic = {"unstart": "01", "running": "02", "waiting": "03", "pasused": "04", "finished": "06"}
        args = {'taskId': run_id, 'token': web_config.token, "statusid": status_dic[status],
                "ostatusid": status_dic[status], "bianm": web_config.bianma}
        url = web_config.status_change_url
        request = GuanHTTP()
        try:
            response = request.post(url=url, json=args)
            if response["res"] == 0:
                self.insert_log(status=False, task_id="", operation="push status change", information=str(response["value"]))
                logger.error(str(response["value"]))
            else:
                self.insert_log(status=True, task_id="", operation="push status change", information="success")
                logger.info("push status change success")
        except Exception as e:
            print("can't reach status_change_url")

    def get_task_result_total_page(self, p):
        task_list_db = self._m_db["task_list"]
        ret = 0
        page_size = 11
        try:
            p['run_id']
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}
        run_id = p['run_id']
        try:
            if task_list_db.count_documents({"run_id": run_id}) > 0:
                pass
            else:
                return {"res": 0, "value": "no such task"}
            self._m_db_collection_cases = self._m_db[run_id + "-cases"]
            c = self._m_db_collection_cases.count_documents({})
            ret = int(c / page_size) + 1
            # return {"res": 1, "value": ret}
            return {"res": 1, "value": c}
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}

    def log(self, p):
        log = self._m_db["log"]
        try:
            p['times']
        except Exception as e:
            self.print_trace_back()
            self.insert_log(status=False, task_id="", operation="get log", information="Lacking parameter " + str(e))
            return {"code": 400, "msg": "Lacking parameter " + str(e), "token": web_config.token}
        start_time = datetime.datetime.strptime(p['times'].split(',')[0], '%Y%m%d')
        end_time = datetime.datetime.strptime(p['times'].split(',')[1], '%Y%m%d')
        log_list = log.find({'time': {'$gt': start_time, '$lt': end_time}}, {"_id": 0, 'time': 0})
        log_result = []
        for log_info in log_list:
            log_result.append(log_info)
        self.insert_log(status=True, task_id="", operation="get log", information="success")
        return {"code": 200, "msg": "success", "data": log_result, "token": web_config.token}

    def get_task_result_by_page(self, p, page_id):
        task_list_db = self._m_db["task_list"]
        ret = []
        page_size = 1
        try:
            p['run_id']
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}
        run_id = p["run_id"]
        try:
            if page_id <= 0:
                return {"res": 0, "value": ret}
            if task_list_db.count_documents({"run_id": run_id}) > 0:
                pass
            else:
                return {"res": 0, "value": "no such task"}
            self._m_db_collection_cases = self._m_db[run_id + "-cases"]
            self._m_db_collection_steps = self._m_db[run_id + "-steps"]
            l = self._m_db_collection_cases.count_documents({})
            page_max = int(l / page_size) + 1
            if page_id < page_max:
                ret_all = self._m_db_collection_steps.find({},
                                                           {"_id": 0, "test_case_index": 1, "type": 1, "description": 1,
                                                            "data": 1,
                                                            "timestamp2": 1, "istrancated": 1}).limit(
                    page_size * 8).skip((page_id - 1) * page_size * 8)
            else:
                ret_all = self._m_db_collection_steps.find({},
                                                           {"_id": 0, "test_case_index": 1, "type": 1, "description": 1,
                                                            "data": 1,
                                                            "timestamp2": 1, "istrancated": 1}).limit(
                    (l - (page_max - 1) * page_size) * 8).skip((page_max - 1) * page_size * 8)

            if ret_all is not None:
                res_list = []
                row = next(ret_all)
                flag = row["test_case_index"]
                res_list.append(row)
                for row in ret_all:
                    if flag == row["test_case_index"]:
                        res_list.append(row)
                    else:
                        ret.append(res_list)
                        flag = row["test_case_index"]
                        res_list = []
                        res_list.append(row)
                ret.append(res_list)
            return {"res": 1, "value": ret}
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}

    def get_task_crash_total_page(self, p):
        task_list_db = self._m_db["task_list"]
        page_size = 3
        try:
            p['run_id']
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}
        run_id = p["run_id"]
        try:
            if task_list_db.count_documents({"run_id": run_id}) > 0:
                pass
            else:
                return {"res": 0, "value": "no such task"}
            self._m_db_collection_crash_bin = self._m_db[run_id + "-crash_bin"]
            crash_count = self._m_db_collection_crash_bin.count_documents({})
            if crash_count:
                page_max = math.ceil(crash_count / page_size)
                # return {"res": 1, "value": page_max}
                return {"res": 1, "value": crash_count}
            else:
                return {"res": 0, "value": "no crash"}
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}

    def get_task_crash_by_page(self, p, page_id):
        task_list_db = self._m_db["task_list"]
        ret = []
        page_size = 3
        try:
            p['run_id']
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}
        run_id = p["run_id"]
        try:
            if page_id <= 0:
                return {"res": 0, "value": ret}
            if task_list_db.count_documents({"run_id": run_id}) > 0:
                pass
            else:
                return {"res": 0, "value": "no such task"}
            self._m_db_collection_crash_bin = self._m_db[run_id + "-crash_bin"]
            crash_count = self._m_db_collection_crash_bin.count_documents({})
            if not crash_count:
                return {"res": 0, "value": "no crash"}
            page_max = math.ceil(crash_count / page_size)
            if page_id < page_max:
                ret_all = self._m_db_collection_crash_bin.find({}, {"_id": 0}).limit(page_size).skip(
                    (page_id - 1) * page_size)
            else:
                ret_all = self._m_db_collection_crash_bin.find({}, {"_id": 0}).skip((page_max - 1) * page_size)
            return {"res": 1, "value": list(ret_all)}
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}


    def delete_task(self, p):
        task_list_db = self._m_db["task_list"]
        try:
            p['run_id']
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}
        run_id = p["run_id"]
        try:
            if task_list_db.count_documents({"run_id": run_id}) > 0:
                pass
            else:
                return {"res": 0, "value": "no such task"}

            self._m_db_collection_cases = self._m_db[run_id + "-cases"]
            self._m_db_collection_cases.drop()
            self._m_db_collection_steps = self._m_db[run_id + "-steps"]
            self._m_db_collection_steps.drop()
            self._m_db_collection_crash_bin = self._m_db[run_id + "-crash_bin"]
            self._m_db_collection_crash_bin.drop()
            mclient = self._db_handler.getMC()
            mdb = mclient["dongjianpcap"]
            chunks = mdb[run_id + ".chunks"]
            fies = mdb[run_id + ".files"]
            chunks.drop()
            fies.drop()
            seed = mclient["seed"]
            seed_chunks = seed[run_id + ".chunks"]
            seed_files = seed[run_id + ".files"]
            seed_chunks.drop()
            seed_files.drop()
            task_list_db.delete_one({"run_id": run_id})
            return {"res": 1, "value": run_id}
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}


    def delete_tasks(self, p):
        try:
            p['taskIds']
        except Exception as e:
            self.print_trace_back()
            self.insert_log(status=False, task_id="", operation="delete task", information="Lacking parameter " + str(e))
            return {"code": 400, "msg": "Lacking parameter " + str(e), "token": web_config.token}
        flag = 1
        for task_id in p['taskIds']:
            if self.delete_task({"run_id": task_id})["res"] == 0:
                flag = 0
                self.insert_log(status=False, task_id=task_id, operation="delete task", information="Something wrong")
            self.insert_log(status=True, task_id=task_id, operation="delete task", information="success")
        if flag == 1:
            return {"code": 200, "msg": "success", "token": web_config.token}
        return {"code": 400, "msg": "Failed to delete some tasks.", "token": web_config.token}

    def get_task_parameter(self, p):
        task_list_db = self._m_db["task_list"]
        try:
            p['run_id']
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}
        run_id = p["run_id"]
        try:
            parameter = task_list_db.find({"run_id": run_id}, {"_id": 0})
            parameter_list = parameter.next()["params"]
            if parameter_list:
                return {"res": 1, "value": parameter_list}
            else:
                return {"res": 0, "value": "no such task"}
        except Exception as e:
            self.print_trace_back()
            return {"res": 0, "value": str(e)}

    def error_log(self, err):
        log = open("err_log.log", "a+")
        try:
            time = datetime.datetime.now().isoformat().replace(":", "-").replace(".", "-")
            log.write(time + ":\r\n")
            log.write(err)
        except Exception as e:
            traceback.print_exc()
        finally:
            log.close()

    def print_trace_back(self):
        if self.debug:
            traceback.print_exc()
            self.error_log(traceback.format_exc())

    def little(self, a, b):
        if a < b:
            return a
        return b

    def transtomac(self, mac):
        sss = 0
        for s in range(0, 11, 2):
            ss = self.TransToInt(mac[s]) * 16 + self.TransToInt(mac[s + 1])
            sss = sss * 256 + ss
        mac2 = sss.to_bytes(6, byteorder="big", signed=False)
        return mac2

    def TransToInt(self, c):  # trans a char to a int
        if '0' <= c <= '9':
            return int(c)
        if 'a' <= c <= 'z':
            return ord(c) - ord('a') + 10
        return 0

    def get_file_name(self, p):
        file_name = None
        try:
            p['run_id']
        except Exception as e:
            return file_name
        run_id = p["run_id"]
        file_name = run_id
        return file_name

    # def push_backup(self, currenttime):
    #     # guan kong platform backup url
    #     BACKUP_PATH = "../backup"
    #     try:
    #         url = web_config.backup_push_url
    #         uploadhttp = GuanHTTP()
    #         files = {
    #             "file": (currenttime + ".tar.gz", open(BACKUP_PATH + "/" + currenttime + ".tar.gz", "rb"))
    #         }
    #         args = {"id": id, "token": web_config.token}
    #         response = uploadhttp.post(url=url, files=files, json=args)
    #         if response["res"] == 0:
    #             self.insert_log(status=False, task_id="", operation="backup data", information=str(response["value"]))
    #             logger.error(str(response["value"]))
    #         else:
    #             self.insert_log(status=True, task_id="", operation="backup data", information="upload file success")
    #             logger.info("upload file success")
    #     except Exception as e:
    #         print("can't reach result_push_url")

    def push_backup(self, id, status, type, msg, data):
        args = {'status': status, 'token': web_config.token, "type": type, 'id': id, "bianm": web_config.bianma, "msg": msg, "data": data}
        url = web_config.backup_push_url
        request = GuanHTTP()
        try:
            response = request.post(url=url, json=args)
            if response["res"] == 0:
                self.insert_log(status=False, task_id="", operation="push backup result",
                                information=str(response["value"]))
                logger.error(str(response["value"]))
            else:
                self.insert_log(status=True, task_id="", operation="push backup result", information="success")
                logger.info("push backup success")
        except Exception as e:
            print("can't reach backup_push_url")

    def backup_data(self, p):
        try:
            p['token']
        except Exception as e:
            self.print_trace_back()
            self.insert_log(status=False, task_id="", operation="backup data", information="Lacking parameter " + str(e))
            return {"code": 400, "msg": "Lacking parameter " + str(e), "token": web_config.token}
        currenttime = datetime.datetime.now().isoformat().replace(":", "-").replace(".", "-")
        BACKUP_PATH = "../backup"
        if not os.path.exists(BACKUP_PATH):
            os.makedirs(BACKUP_PATH + "/" + currenttime)

        try:
            # backup mongo
            dump_mongo_cmd = web_config.mongodb_location + "mongodump -h {} -u {} -p {} -d {} --authenticationDatabase admin -o {}/{}/mongo 2>/dev/null".format(
                db_config.db_host,
                db_config.db_user,
                db_config.db_passwd,
                db_config.db_name,
                BACKUP_PATH,
                currenttime)
            os.system(dump_mongo_cmd)

            compress_cmd = "cd {0};/bin/tar czf {1}.tar.gz {1}".format(BACKUP_PATH, currenttime)
            os.system(compress_cmd)
            logger.info("local database backup success.")

            self.push_backup(currenttime, 1, 1, "backup success", currenttime)
            # push_backup_thread = threading.Thread(target=self.push_backup, args=[currenttime,])
            # push_backup_thread.start()
            return ({"code": 200, "msg": "backup success", "id": currenttime + ".tar.gz", "token": web_config.token})
        except Exception as e:
            self.push_backup(currenttime, 0, 1, "backup failed: " + str(e), currenttime)
            return ({"code": 400, "msg": "backup failed: " + str(e), "id": currenttime + ".tar.gz", "token": web_config.token})

        # shutil.rmtree(BACKUP_PATH)

    def restore_data(self, p):
        try:
            p['id']
        except Exception as e:
            self.print_trace_back()
            self.insert_log(status=False, task_id="", operation="backup data", information="Lacking parameter " + str(e))
            return {"code": 400, "msg": "Lacking parameter " + str(e), "token": web_config.token}

        RESTORE_PATH = "../backup"
        file_path = RESTORE_PATH + "/" + p['id'].split('.tar.gz')[0]

        if not os.path.exists(file_path):
            return {"code": 400, "msg": file_path + "is not exit, please backup first", "token": web_config.token}

        # url = "http://192.168.182.230:888/downloader/%s" % data_date
        # downloadhttp = GuanHTTP()
        # response = downloadhttp.getFile(url, "{}/{}.tar.gz".format(RESTORE_PATH, data_date))
        # if response["res"] == 0:
        #     self.insert_log(task_id="", operation="retore data", information=str(response["value"]))
        #     logger.error(response["value"])

        # decompress
        # decompress_cmd = "cd {};/bin/tar xzf {}".format(RESTORE_PATH, p['id'])
        # os.system(decompress_cmd)

        # load_mysql_cmd = "cd {};/usr/bin/mysql -h {} -u {} -p'{}' {} < {}/mysql 2>/dev/null".format(
        #     RESTORE_PATH,
        #     db_config.db_host,
        #     db_config.db_user,
        #     db_config.db_passwd,
        #     db_config.db_name,
        #     data_date)
        # os.system(load_mysql_cmd)
        try:
            load_mongo_cmd = 'ulimit -n 65535;' + web_config.mongodb_location + "mongorestore -h {} -u {} -p {} -d {} --drop --authenticationDatabase admin {}/mongo/{}".format(
                db_config.db_host,
                db_config.db_user,
                db_config.db_passwd,
                db_config.db_name,
                file_path,
                db_config.db_name
            )
            os.system(load_mongo_cmd)
            self.insert_log(status=True, task_id="", operation="retore data", information="success")
            logger.info("local database restore success.")
            self.push_backup(p['id'], 1, 0, "restore success", p['id'])
            return {"code": 200, "msg": "restore success", "token": web_config.token}
        except Exception as e:
            self.push_backup(p['id'], 0, 0, "restore success", p['id'])
            return {"code": 400, "msg": "restore failed: " + str(e), "token": web_config.token}
        # shutil.rmtree(RESTORE_PATH)

    def gen_pcap(self, run_id):
        mclient = self._db_handler.getMC()
        mdb = mclient["dongjianpcap"]
        try:
            TMP_PATH = "../tmp"
            if not os.path.exists(TMP_PATH + "/" + run_id):
                os.makedirs(TMP_PATH + "/" + run_id)
            buket = GridFSBucket(mdb, bucket_name=run_id)
            for file in buket.find():
                pcapfile = TMP_PATH + "/" + run_id + "/" + file.filename
                with open(pcapfile, "wb") as f:
                    buket.download_to_stream_by_name(file.filename, f)
            compress_cmd = "cd {0};/bin/tar czf {1}.tar.gz {1}".format(TMP_PATH, run_id)
            os.system(compress_cmd)
            shutil.rmtree(TMP_PATH + "/" + run_id)
            logger.info("{} pcap generate success.".format(run_id))
            return {"res": 1, "value": run_id}
        except Exception as e:
            self.print_trace_back()
            logger.error(str(e))

    def replay_pcap(self, run_id, case_id):
        TMP_PATH = "../pcaptmp"
        if not os.path.exists(TMP_PATH):
            os.mkdir(TMP_PATH)
        mclient = self._db_handler.getMC()
        mdb = mclient["dongjianpcap"]
        buket = GridFSBucket(mdb, bucket_name=run_id)

        mpcap = buket.find({"filename": "{}-{}.pcap".format(run_id, case_id)}, no_cursor_timeout=True)
        if mpcap is None:
            return {"res": 0, "value": "run_id or case_id may be wrong."}
        for file in mpcap:  # size default 1
            pcapfile = TMP_PATH + "/" + file.filename
            with open(pcapfile, "wb") as f:
                buket.download_to_stream_by_name(file.filename, f)
            packets = rdpcap("{}/{}-{}.pcap".format(TMP_PATH, run_id, case_id))
            # check packet layer
            try:
                packets[0]['TCP']
                packettype = 0
            except IndexError:
                try:
                    packets[0]['UDP']
                    packettype = 1
                except IndexError:
                    try:
                        packets[0]['IP']
                        packettype = 3  # law-3
                    except IndexError:
                        packettype = 2  # law-2
            if packettype == 2:
                sendp(packets)
            elif packettype == 3:
                sendp(packets)
            elif packettype == 1:
                sendp(packets)
            else:
                try:
                    packets[0]['TCP']
                    packettype = 0
                except IndexError:
                    try:
                        packets[0]['UDP']
                        packettype = 1
                    except IndexError:
                        try:
                            packets[0]['IP']
                            packettype = 3  # law-3
                        except IndexError:
                            packettype = 2  # law-2
                if packettype == 2:
                    sendp(packets)
                elif packettype == 3:
                    sendp(packets)
                elif packettype == 1:
                    sendp(packets)
                else:
                    def gentcpdata(packets):
                        data = b''
                        for packet in packets:
                            if packet["TCP"].flags == "PA":
                                data = data + packet["TCP"].load
                            elif packet["TCP"].flags == "FPA":
                                data = data + packet["TCP"].load
                                yield data
                                data = b''
                            else:  # small packet, no need Reorganization
                                data = data + packet["TCP"].load
                                yield data
                                data = b''

                    _sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                    _sock.connect((packets[0]['IP'].dst, packets[0]['TCP'].dport))
                    datas = gentcpdata(packets)
                    # no consider recv
                    for data in datas:
                        _sock.send(data)
                    _sock.close()
        shutil.rmtree(TMP_PATH)
        return {"res": 1, "value": "{}-{} send success.".format(run_id, case_id)}

    def export_document(self, document):
        run_id, doc_format = document.split('.')
        if os.path.exists("../" + doc_format + "/" + run_id + "." + doc_format):
            return r"../" + doc_format + "/"
        mclient = self._db_handler.getMC()
        mdb = mclient["report"]
        report = GridFSBucket(mdb, bucket_name=run_id + "-report")
        fp = open(r'../' + doc_format + '/' + run_id + '.' + doc_format, 'w')
        report.download_to_stream(run_id + "." + doc_format, fp)
        fp.close()
        return r"../" + doc_format + "/"

    # def get_netzob_list_total_page(self):
    #     self.func_desc = "get_netzob_list_total_page: "
    #     res = 0
    #     page_size = 11
    #
    #     try:
    #         c = self._m_mycol_netzob.find({}).count()
    #         print(c)
    #         if c > 0:
    #             if c % page_size == 0:
    #                 res=int(c / page_size)
    #             else:
    #                 res = int(c / page_size) + 1
    #         return {"res": 1, "value": res}
    #     except Exception as e:
    #         print(self.func_desc + str(e))
    #         self.print_trace_back()
    #         return {"res": 0, "value": str(e)}
    #
    #
    # def get_netzob_list_by_page(self, page_id):
    #     self._conn_r = self._db_handler.getCon()
    #     self._cursor_r = self._conn_r.cursor()
    #     self.func_desc = "get_netzob_list_by_page: "
    #     res = []
    #     page_size = 11
    #     collections=[]
    #     try:
    #         c = int(self._m_mycol_netzob.find({}).count())
    #         page_max = int(c/ page_size) + 1
    #         if page_id <= 0:
    #             return {"res": 0, "value": res}
    #         if page_id < page_max:
    #             collections = self._m_mycol_netzob.find({}, {"_id": 0}).limit(page_size).skip((page_id-1)*page_size)
    #         else:
    #
    #             collections = self._m_mycol_netzob.find({}, {"_id": 0}).skip((page_max-1)*page_size)
    #         if c > 0:
    #             for item in collections:
    #                 res.append(item)
    #         return {"res": 1, "value": res}
    #     except Exception as e:
    #         print(self.func_desc + str(e))
    #         self.print_trace_back()
    #         return {"res": 0, "value": str(e)}
    #
    #
    #
    # def get_netzob_info(self, p):
    #     self.func_desc = "get_netzob_info: "
    #     try:
    #         p['task_name']
    #     except Exception as e:
    #         print(self.func_desc + "1" + str(e))
    #         self.print_trace_back()
    #         return {"res": 0, "value": str(e)}
    #     try:
    #         ret_all = self._m_mycol_netzob.find_one({'task_name': p['task_name']},{"_id": 0})
    #         return {"res": 1, "value": ret_all}
    #     except Exception as e:
    #         print(self.func_desc + "-2 " + str(e))
    #         self.print_trace_back()
    #         return {"res": 0, "value": str(e)}
